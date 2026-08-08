"""WP-R1: native MS Office ingest without Docling."""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from unittest.mock import MagicMock

from aerobim.application.services.package_ingestion import PackageIngestionService
from aerobim.domain.models import (
    CapabilityState,
    RequirementSource,
    SourceKind,
    ValidationRequest,
)
from aerobim.infrastructure.adapters.docling_requirement_extractor import (
    StructuredRequirementExtractor,
)
from aerobim.infrastructure.adapters.narrative_rule_synthesizer import (
    NarrativeRuleSynthesizer,
)
from aerobim.infrastructure.adapters.docling_office_document_ingestor import (
    DoclingOfficeDocumentIngestor,
    OfficeDocumentIngestor,
    _LEGACY_FAIL_MESSAGE,
)


class OfficeNativeIngestTests(unittest.TestCase):
    def test_docx_ingest_without_docling(self) -> None:
        try:
            from docx import Document
        except ModuleNotFoundError:
            self.skipTest("python-docx not installed")

        ingestor = OfficeDocumentIngestor()
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "memo.docx"
            doc = Document()
            doc.add_paragraph("Пояснительная записка: площадь помещения 120.5 м2")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Параметр"
            table.cell(0, 1).text = "Значение"
            table.cell(1, 0).text = "Площадь"
            table.cell(1, 1).text = "120.5"
            doc.save(path)
            source = ingestor.ingest(path)
        self.assertIn("120.5", source.text)
        self.assertIn("provenance: paragraph:1", source.text)
        self.assertIn("table:1!R2C2", source.text)
        self.assertEqual(source.source_kind, SourceKind.TECHNICAL_SPECIFICATION)

    def test_xlsx_ingest_with_cell_provenance(self) -> None:
        try:
            from openpyxl import Workbook
        except ModuleNotFoundError:
            self.skipTest("openpyxl not installed")

        ingestor = OfficeDocumentIngestor()
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "spec.xlsx"
            wb = Workbook()
            ws = wb.active
            assert ws is not None
            ws.title = "Areas"
            ws["A1"] = "rule"
            ws["B1"] = "IfcSpace"
            ws["A2"] = "R-AREA"
            ws["B2"] = "120.5"
            wb.save(path)
            source = ingestor.ingest(path)
        self.assertIn("sheet!Areas!R2C2", source.text)
        self.assertIn("120.5", source.text)

    def test_legacy_doc_fail_closed(self) -> None:
        ingestor = OfficeDocumentIngestor()
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "legacy.doc"
            path.write_bytes(b"\xd0\xcf\x11\xe0")
            with self.assertRaisesRegex(RuntimeError, "available_as_derived_input"):
                ingestor.ingest(path)

    def test_hydrate_sets_office_capability_ok(self) -> None:
        try:
            from openpyxl import Workbook
        except ModuleNotFoundError:
            self.skipTest("openpyxl not installed")

        service = PackageIngestionService(
            drawing_analyzer=MagicMock(),
            narrative_rule_synthesizer=NarrativeRuleSynthesizer(),
            office_document_ingestor=OfficeDocumentIngestor(),
        )
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "spec.xlsx"
            wb = Workbook()
            ws = wb.active
            assert ws is not None
            ws["A1"] = "R1|IfcSpace|Pset_SpaceCommon|NetPlannedArea|eq|120.5"
            wb.save(path)
            request = ValidationRequest(
                request_id="req-1",
                ifc_path=Path(tmp) / "m.ifc",
                requirement_source=RequirementSource(path=path),
            )
            hydrated, cap = service.maybe_hydrate_office_requirement_source(request)
        self.assertIsNotNone(cap)
        assert cap is not None
        self.assertEqual(cap.status, CapabilityState.OK)
        self.assertIn("office_ingest:xlsx", cap.reason or "")
        self.assertIn("120.5", hydrated.requirement_source.text)

    def test_cross_doc_pipe_row_extracted_from_xlsx(self) -> None:
        try:
            from openpyxl import Workbook
        except ModuleNotFoundError:
            self.skipTest("openpyxl not installed")

        extractor = StructuredRequirementExtractor()
        ingestor = OfficeDocumentIngestor()
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "spec.xlsx"
            wb = Workbook()
            ws = wb.active
            assert ws is not None
            ws["A1"] = "R-AREA|IFCWALL|Pset_SpaceCommon|NetPlannedArea|120.5"
            wb.save(path)
            source = ingestor.ingest(path)
            reqs = list(extractor.extract(source))
        self.assertEqual(len(reqs), 1)
        self.assertEqual(reqs[0].rule_id, "R-AREA")
        self.assertIn("120.5", reqs[0].expected_value or "")


if __name__ == "__main__":
    unittest.main()
