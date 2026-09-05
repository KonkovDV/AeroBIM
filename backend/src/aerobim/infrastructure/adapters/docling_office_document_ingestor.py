"""OfficeDocumentIngestor — native python-docx/openpyxl by default; Docling optional."""

from __future__ import annotations

import re
from pathlib import Path

from aerobim.domain.models import RequirementSource, SourceKind

_TEXT_SUFFIXES = {".txt", ".md", ".csv"}
_NATIVE_DOCX = {".docx"}
_NATIVE_XLSX = {".xlsx"}
_LEGACY_BINARY = {".doc", ".xls"}
_DOCLING_SUFFIXES = {".pptx", ".ppt", ".odt", ".ods"}
_OFFICE_SUFFIXES = _NATIVE_DOCX | _NATIVE_XLSX | _LEGACY_BINARY | _DOCLING_SUFFIXES

_MAX_OFFICE_BYTES = 25 * 1024 * 1024
_MAX_ZIP_MEMBERS = 512
_LEGACY_FAIL_MESSAGE = (
    "Legacy binary Office format not supported; convert to OOXML (.docx/.xlsx/.pptx). "
    "available_as_derived_input"
)


class OfficeDocumentIngestor:
    """Extract text from MS Office documents into RequirementSource with cell/para provenance."""

    def ingest(self, path: Path) -> RequirementSource:
        if not path.exists():
            raise FileNotFoundError(path)
        if path.stat().st_size > _MAX_OFFICE_BYTES:
            raise ValueError(f"Office file exceeds {_MAX_OFFICE_BYTES} bytes: {path.name}")

        suffix = path.suffix.lower()
        if suffix in _LEGACY_BINARY:
            raise RuntimeError(_LEGACY_FAIL_MESSAGE)

        text = self._load_text(path, suffix)
        kind = (
            SourceKind.TECHNICAL_SPECIFICATION
            if suffix in _OFFICE_SUFFIXES
            else SourceKind.STRUCTURED_TEXT
        )
        return RequirementSource(
            text=text,
            path=path,
            source_kind=kind,
            source_id=path.stem,
            doc_type=suffix.lstrip(".") or None,
        )

    def _load_text(self, path: Path, suffix: str) -> str:
        if suffix in _TEXT_SUFFIXES:
            return path.read_text(encoding="utf-8")
        if suffix in _NATIVE_DOCX:
            return self._ingest_docx(path)
        if suffix in _NATIVE_XLSX:
            return self._ingest_xlsx(path)
        if suffix == ".pptx":
            self._guard_zip_bomb(path)
        return self._ingest_via_docling(path, suffix)

    def _guard_zip_bomb(self, path: Path) -> None:
        if path.suffix.lower() not in {".docx", ".xlsx", ".pptx"}:
            return
        from aerobim.core.security.zip_limits import (
            ZipBombError,
            inspect_zip_path,
            verify_zip_inflate,
        )

        try:
            inspect_zip_path(path, max_members=_MAX_ZIP_MEMBERS)
            verify_zip_inflate(path)
        except ZipBombError as exc:
            raise ValueError(str(exc)) from exc

    def _ingest_docx(self, path: Path) -> str:
        self._guard_zip_bomb(path)
        try:
            from docx import Document
        except ModuleNotFoundError as exc:
            raise RuntimeError(
                "python-docx is required for .docx ingest (core dependency)"
            ) from exc
        document = Document(str(path))
        lines: list[str] = []
        for para_idx, paragraph in enumerate(document.paragraphs, start=1):
            text = paragraph.text.strip()
            if text:
                lines.append(f"# provenance: paragraph:{para_idx}")
                lines.append(text)
        for table_idx, table in enumerate(document.tables, start=1):
            for row_idx, row in enumerate(table.rows, start=1):
                for col_idx, cell in enumerate(row.cells, start=1):
                    cell_text = cell.text.strip()
                    if cell_text:
                        prov = f"table:{table_idx}!R{row_idx}C{col_idx}"
                        lines.append(f"# provenance: {prov}")
                        lines.append(cell_text)
        return "\n".join(lines).strip()

    def _ingest_xlsx(self, path: Path) -> str:
        self._guard_zip_bomb(path)
        try:
            from openpyxl import load_workbook
        except ModuleNotFoundError as exc:
            raise RuntimeError("openpyxl is required for .xlsx ingest (core dependency)") from exc
        workbook = load_workbook(path, read_only=True, data_only=True)
        lines: list[str] = []
        try:
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                    for col_idx, value in enumerate(row, start=1):
                        if value is None or str(value).strip() == "":
                            continue
                        prov = f"sheet!{sheet_name}!R{row_idx}C{col_idx}"
                        cell_text = str(value).strip()
                        lines.append(f"# provenance: {prov}")
                        lines.append(cell_text)
        finally:
            workbook.close()
        return "\n".join(lines).strip()

    def _ingest_via_docling(self, path: Path, suffix: str) -> str:
        try:
            from docling.document_converter import DocumentConverter
        except ModuleNotFoundError as exc:
            raise RuntimeError(
                f"Docling optional extra required for {suffix} "
                "(pip install aerobim-backend[docling]); "
                f"native path supports .docx/.xlsx only"
            ) from exc
        converter = DocumentConverter()
        result = converter.convert(str(path))
        return self._normalize_markdown(result.document.export_to_markdown())

    def _normalize_markdown(self, markdown: str) -> str:
        normalized_lines: list[str] = []
        for line in markdown.splitlines():
            stripped = line.strip()
            if not stripped:
                normalized_lines.append("")
                continue
            stripped = re.sub(r"^#+\s*", "", stripped)
            stripped = re.sub(r"^[-*+]\s*", "", stripped)
            stripped = stripped.replace(r"\_", "_")
            normalized_lines.append(stripped)
        return "\n".join(normalized_lines)


DoclingOfficeDocumentIngestor = OfficeDocumentIngestor

__all__ = ["DoclingOfficeDocumentIngestor", "OfficeDocumentIngestor", "_LEGACY_FAIL_MESSAGE"]
