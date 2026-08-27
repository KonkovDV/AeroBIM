"""Extract declared calc/BIM tables from xlsx/docx. Native .lir and PDF stay closed."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Literal

from aerobim.domain.calculation_table_compare import (
    DeclaredCalcRow,
    TableCompareResult,
    compare_declared_tables,
)

ExtractStatus = Literal[
    "ok",
    "native_lir_not_implemented",
    "pdf_fragile",
    "empty",
    "unreadable",
    "unsupported",
]

_NATIVE_LIR = frozenset({".lir", ".spr"})
_PDF = frozenset({".pdf"})
_XLSX = frozenset({".xlsx", ".xlsm"})
_DOCX = frozenset({".docx"})

_ID_HEADERS = frozenset({"id", "field_id", "load_id", "field"})
_LABEL_HEADERS = frozenset({"label", "name", "title"})
_VALUE_HEADERS = frozenset({"value", "declared", "amount"})
_EXPECTED_HEADERS = frozenset({"expected"})
_OBSERVED_HEADERS = frozenset({"observed"})
_UNIT_HEADERS = frozenset({"unit", "ед", "ед."})


@dataclass(frozen=True)
class ExtractedTable:
    status: ExtractStatus
    rows: tuple[DeclaredCalcRow, ...]
    path_suffix: str
    reason: str


def extract_declared_rows(path: Path) -> ExtractedTable:
    """Read a declared-field table. Never parses native LIRA binaries or PDF tables."""

    suffix = path.suffix.lower()
    if suffix in _NATIVE_LIR:
        return ExtractedTable(
            status="native_lir_not_implemented",
            rows=(),
            path_suffix=suffix,
            reason="native .lir/.spr is not implemented",
        )
    if suffix in _PDF:
        return ExtractedTable(
            status="pdf_fragile",
            rows=(),
            path_suffix=suffix,
            reason="PDF table compare remains fragile; use xlsx/docx",
        )
    if not path.is_file():
        return ExtractedTable(
            status="unreadable",
            rows=(),
            path_suffix=suffix,
            reason="path is not a file",
        )
    if suffix in _XLSX:
        return _extract_xlsx(path, suffix)
    if suffix in _DOCX:
        return _extract_docx(path, suffix)
    return ExtractedTable(
        status="unsupported",
        rows=(),
        path_suffix=suffix,
        reason=f"suffix {suffix} is not an office declared-table source",
    )


def compare_office_tables(calc_path: Path, bim_path: Path) -> TableCompareResult:
    """Two-file office compare. Callers must not treat MATCH as solver OK."""

    calc = extract_declared_rows(calc_path)
    bim = extract_declared_rows(bim_path)
    return compare_declared_tables(calc.rows, bim.rows)


def _norm_header(raw: object) -> str:
    return str(raw or "").strip().casefold()


def _header_map(cells: list[str]) -> dict[str, int] | None:
    mapping: dict[str, int] = {}
    for index, cell in enumerate(cells):
        token = _norm_header(cell)
        if token in _ID_HEADERS:
            mapping["id"] = index
        elif token in _LABEL_HEADERS:
            mapping["label"] = index
        elif token in _VALUE_HEADERS:
            mapping["value"] = index
        elif token in _EXPECTED_HEADERS:
            mapping["expected"] = index
        elif token in _OBSERVED_HEADERS:
            mapping["observed"] = index
        elif token in _UNIT_HEADERS:
            mapping["unit"] = index
    if "id" not in mapping:
        return None
    if "value" not in mapping and "expected" not in mapping:
        return None
    return mapping


def _row_from_cells(cells: list[str], mapping: dict[str, int]) -> DeclaredCalcRow | None:
    field_id = cells[mapping["id"]].strip() if mapping["id"] < len(cells) else ""
    if not field_id:
        return None
    label = ""
    if "label" in mapping and mapping["label"] < len(cells):
        label = cells[mapping["label"]].strip()
    unit = ""
    if "unit" in mapping and mapping["unit"] < len(cells):
        unit = cells[mapping["unit"]].strip()
    value = ""
    if "value" in mapping and mapping["value"] < len(cells):
        value = cells[mapping["value"]].strip()
    elif "expected" in mapping and mapping["expected"] < len(cells):
        value = cells[mapping["expected"]].strip()
    if not value:
        return None
    return DeclaredCalcRow(field_id=field_id, label=label, value=value, unit=unit)


def _extract_xlsx(path: Path, suffix: str) -> ExtractedTable:
    try:
        from openpyxl import load_workbook
    except ModuleNotFoundError:
        return ExtractedTable(
            status="unreadable",
            rows=(),
            path_suffix=suffix,
            reason="openpyxl is required for .xlsx declared tables",
        )
    workbook = load_workbook(path, read_only=True, data_only=True)
    try:
        sheet = workbook[workbook.sheetnames[0]]
        matrix: list[list[str]] = []
        for row in sheet.iter_rows(values_only=True):
            matrix.append(["" if cell is None else str(cell).strip() for cell in row])
    finally:
        workbook.close()
    return _rows_from_matrix(matrix, suffix)


def _extract_docx(path: Path, suffix: str) -> ExtractedTable:
    try:
        from docx import Document
    except ModuleNotFoundError:
        return ExtractedTable(
            status="unreadable",
            rows=(),
            path_suffix=suffix,
            reason="python-docx is required for .docx declared tables",
        )
    document = Document(str(path))
    if not document.tables:
        return ExtractedTable(
            status="empty",
            rows=(),
            path_suffix=suffix,
            reason="docx has no tables",
        )
    table = document.tables[0]
    matrix = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    return _rows_from_matrix(matrix, suffix)


def _rows_from_matrix(matrix: list[list[str]], suffix: str) -> ExtractedTable:
    mapping: dict[str, int] | None = None
    start = 0
    for index, row in enumerate(matrix):
        found = _header_map(row)
        if found is not None:
            mapping = found
            start = index + 1
            break
    if mapping is None:
        return ExtractedTable(
            status="empty",
            rows=(),
            path_suffix=suffix,
            reason="no id+value/expected header row",
        )
    rows: list[DeclaredCalcRow] = []
    for row in matrix[start:]:
        parsed = _row_from_cells(row, mapping)
        if parsed is not None:
            rows.append(parsed)
    if not rows:
        return ExtractedTable(
            status="empty",
            rows=(),
            path_suffix=suffix,
            reason="header found but no data rows",
        )
    return ExtractedTable(status="ok", rows=tuple(rows), path_suffix=suffix, reason="ok")


__all__ = [
    "ExtractedTable",
    "compare_office_tables",
    "extract_declared_rows",
]
