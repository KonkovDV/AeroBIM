"""Load-table numeric сверка from calculation sources (not correctness)."""

from __future__ import annotations

import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Literal

from aerobim.domain.calculation_table_compare import (
    DeclaredCalcRow,
    TableCompareResult,
    compare_declared_tables,
)
from aerobim.domain.models import FindingCategory, Severity, ValidationIssue, ValidationRequest
from aerobim.domain.quantity import parse_quantity, si_compare

_LOAD_ROW = re.compile(
    r"(?P<id>[A-Za-zА-Яа-я0-9_.\-]+)\s*[|;]\s*"
    r"(?P<label>[^|;]+)\s*[|;]\s*"
    r"(?P<expected>-?\d+(?:[.,]\d+)?)\s*[|;]\s*"
    r"(?P<unit>[A-Za-zА-Яа-я²³23/]*)\s*[|;]\s*"
    r"(?P<observed>-?\d+(?:[.,]\d+)?)",
    re.IGNORECASE,
)


def _fail_closed_office_or_lira(
    path: Path | None,
    *,
    source_id: str,
) -> list[ValidationIssue] | None:
    """Native LIRA / PDF / office binaries must not be utf-8-decoded as LOAD text."""

    if path is None:
        return None
    suffix = path.suffix.lower()
    if suffix in {".lir", ".spr"}:
        return [
            ValidationIssue(
                rule_id="AEROBIM-LIRA-NATIVE",
                severity=Severity.WARNING,
                message="Native LIRA .lir/.spr is not implemented (not a solver skip)",
                category=FindingCategory.CROSS_DOCUMENT,
                source_id=source_id,
            )
        ]
    if suffix == ".pdf":
        return [
            ValidationIssue(
                rule_id="AEROBIM-LIRA-PDF",
                severity=Severity.INFO,
                message="LIRA PDF table compare remains fragile; use xlsx/docx",
                category=FindingCategory.CROSS_DOCUMENT,
                source_id=source_id,
            )
        ]
    if suffix not in {".xlsx", ".xlsm", ".docx"}:
        return None
    extracted = extract_declared_rows(path)
    if extracted.status != "ok":
        return [
            ValidationIssue(
                rule_id="AEROBIM-LIRA-OFFICE",
                severity=Severity.INFO,
                message=f"Office declared table: {extracted.reason}",
                category=FindingCategory.CROSS_DOCUMENT,
                source_id=source_id,
            )
        ]
    return [
        ValidationIssue(
            rule_id="AEROBIM-LIRA-OFFICE",
            severity=Severity.INFO,
            message=(
                f"Office declared table has {len(extracted.rows)} row(s); "
                "two-file compare is not a solver"
            ),
            category=FindingCategory.CROSS_DOCUMENT,
            source_id=source_id,
        )
    ]


def _numeric_match(expected: float, observed: float, unit: str) -> bool:
    q_e = parse_quantity(expected, unit or "kN")
    q_o = parse_quantity(observed, unit or "kN")
    if q_e.si_value is not None and q_o.si_value is not None:
        return si_compare(q_e, q_o, epsilon=1e-3)
    return abs(expected - observed) <= 1e-3


class SpreadsheetLoadEvidenceAdapter:
    """Parse calculation_source for LOAD|id|label|expected|unit|observed rows or JSON."""

    def verify(self, request: ValidationRequest) -> list[ValidationIssue]:
        source = request.calculation_source
        if source is None:
            return []

        conflict_issues: list[ValidationIssue] = []
        text = ""
        source_id = source.source_id or "calculation"
        closed = _fail_closed_office_or_lira(source.path, source_id=source_id)
        if closed is not None:
            return closed

        # RT-CALC-005: .json path is SSOT when present (text must not shadow).
        if source.path is not None and source.path.suffix.lower() == ".json":
            text = self._load_path(source.path)
            inline = source.text.strip()
            if inline and inline != text.strip():
                conflict_issues.append(
                    ValidationIssue(
                        rule_id="AEROBIM-LOAD-FORMAT",
                        severity=Severity.WARNING,
                        message=(
                            "Calculation source text disagrees with .json path; "
                            "path is SSOT — load OK suppressed"
                        ),
                        category=FindingCategory.CROSS_DOCUMENT,
                        source_id=source_id,
                    )
                )
        else:
            text = source.text.strip()
            if not text and source.path is not None:
                text = self._load_path(source.path)

        if not text.strip():
            return [
                ValidationIssue(
                    rule_id="AEROBIM-LOAD-FORMAT",
                    severity=Severity.INFO,
                    message="Calculation source empty; no LOAD rows evaluated",
                    category=FindingCategory.CROSS_DOCUMENT,
                    source_id=source_id,
                )
            ]

        if text.lstrip().startswith("{"):
            issues = self._verify_json(text, source_id=source_id)
        else:
            issues = self._verify_tabular(text, source_id=source_id)

        if conflict_issues:
            issues = [item for item in issues if item.rule_id != "AEROBIM-LOAD-OK"]
            return [*conflict_issues, *issues]
        return issues

    def _load_path(self, path: Path) -> str:
        if not path.exists():
            raise FileNotFoundError(path)
        return path.read_text(encoding="utf-8")

    def _verify_json(self, text: str, *, source_id: str) -> list[ValidationIssue]:
        try:
            payload = json.loads(text)
        except json.JSONDecodeError as exc:
            return [
                ValidationIssue(
                    rule_id="AEROBIM-LOAD-JSON",
                    severity=Severity.WARNING,
                    message=f"Calculation JSON parse failed: {exc}",
                    category=FindingCategory.CROSS_DOCUMENT,
                    source_id=source_id,
                )
            ]
        rows = payload.get("loads") if isinstance(payload, dict) else None
        if not isinstance(rows, list):
            return [
                ValidationIssue(
                    rule_id="AEROBIM-LOAD-SCHEMA",
                    severity=Severity.INFO,
                    message="Calculation JSON has no 'loads' array; load сверка skipped",
                    category=FindingCategory.CROSS_DOCUMENT,
                    source_id=source_id,
                )
            ]
        issues: list[ValidationIssue] = []
        evaluated_ok = 0
        if len(rows) == 0:
            return [
                ValidationIssue(
                    rule_id="AEROBIM-LOAD-FORMAT",
                    severity=Severity.INFO,
                    message="Calculation JSON 'loads' array is empty; no rows evaluated",
                    category=FindingCategory.CROSS_DOCUMENT,
                    source_id=source_id,
                )
            ]
        for index, row in enumerate(rows):
            # RT-CALC-004: never silently skip non-dict rows.
            if not isinstance(row, dict):
                issues.append(
                    ValidationIssue(
                        rule_id="AEROBIM-LOAD-ROW",
                        severity=Severity.WARNING,
                        message=f"loads[{index}] is not an object",
                        category=FindingCategory.CROSS_DOCUMENT,
                        source_id=source_id,
                    )
                )
                continue
            load_id = str(row.get("id", "load"))
            unit = str(row.get("unit", "") or "")
            try:
                expected = float(row["expected"])
                observed = float(row["observed"])
            except (KeyError, TypeError, ValueError):
                issues.append(
                    ValidationIssue(
                        rule_id="AEROBIM-LOAD-ROW",
                        severity=Severity.WARNING,
                        message=f"Load row {load_id} missing expected/observed numerics",
                        category=FindingCategory.CROSS_DOCUMENT,
                        target_ref=load_id,
                        source_id=source_id,
                    )
                )
                continue
            if not _numeric_match(expected, observed, unit):
                issues.append(
                    ValidationIssue(
                        rule_id="AEROBIM-LOAD-MISMATCH",
                        severity=Severity.WARNING,
                        message=(
                            f"Load match failed for {load_id}: "
                            f"expected={expected} observed={observed} {unit}"
                        ),
                        category=FindingCategory.CROSS_DOCUMENT,
                        target_ref=load_id,
                        expected_value=str(expected),
                        observed_value=str(observed),
                        unit=unit or None,
                        source_id=source_id,
                    )
                )
            else:
                evaluated_ok += 1
        if any(i.rule_id == "AEROBIM-LOAD-MISMATCH" for i in issues):
            return issues
        if any(i.rule_id == "AEROBIM-LOAD-ROW" for i in issues):
            return issues
        if evaluated_ok == 0:
            if not issues:
                return [
                    ValidationIssue(
                        rule_id="AEROBIM-LOAD-FORMAT",
                        severity=Severity.INFO,
                        message="Calculation JSON loads present but no numeric rows evaluated",
                        category=FindingCategory.CROSS_DOCUMENT,
                        source_id=source_id,
                    )
                ]
            return issues
        return [
            ValidationIssue(
                rule_id="AEROBIM-LOAD-OK",
                severity=Severity.INFO,
                message=f"Load сверка matched {evaluated_ok} row(s) (not correctness)",
                category=FindingCategory.CROSS_DOCUMENT,
                source_id=source_id,
            ),
            *issues,
        ]

    def _verify_tabular(self, text: str, *, source_id: str) -> list[ValidationIssue]:
        issues: list[ValidationIssue] = []
        matched_any = False
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped or stripped.startswith("#"):
                continue
            if stripped.upper().startswith("LOAD|") or stripped.upper().startswith("LOAD;"):
                stripped = (
                    stripped.split("|", 1)[-1] if "|" in stripped else stripped.split(";", 1)[-1]
                )
            match = _LOAD_ROW.search(stripped)
            if match is None:
                continue
            matched_any = True
            expected = float(match.group("expected").replace(",", "."))
            observed = float(match.group("observed").replace(",", "."))
            unit = (match.group("unit") or "").strip()
            load_id = match.group("id").strip()
            if not _numeric_match(expected, observed, unit):
                issues.append(
                    ValidationIssue(
                        rule_id="AEROBIM-LOAD-MISMATCH",
                        severity=Severity.WARNING,
                        message=(
                            f"Load match failed for {load_id}: "
                            f"expected={expected} observed={observed} {unit}"
                        ),
                        category=FindingCategory.CROSS_DOCUMENT,
                        target_ref=load_id,
                        expected_value=str(expected),
                        observed_value=str(observed),
                        unit=unit or None,
                        source_id=source_id,
                    )
                )
        if not matched_any:
            return [
                ValidationIssue(
                    rule_id="AEROBIM-LOAD-FORMAT",
                    severity=Severity.INFO,
                    message=(
                        "Calculation source present but no LOAD rows matched; "
                        "expected LOAD|id|label|expected|unit|observed"
                    ),
                    category=FindingCategory.CROSS_DOCUMENT,
                    source_id=source_id,
                )
            ]
        if any(i.rule_id == "AEROBIM-LOAD-MISMATCH" for i in issues):
            return issues
        return [
            ValidationIssue(
                rule_id="AEROBIM-LOAD-OK",
                severity=Severity.INFO,
                message="Load сверка matched tabular row(s) (not correctness)",
                category=FindingCategory.CROSS_DOCUMENT,
                source_id=source_id,
            ),
            *issues,
        ]


ExtractStatus = Literal[
    "ok",
    "native_lir_not_implemented",
    "pdf_fragile",
    "empty",
    "unreadable",
    "unsupported",
]

_NATIVE_LIR = frozenset({".lir", ".spr"})
_PDF = frozenset({".pdf"})
_XLSX = frozenset({".xlsx", ".xlsm"})
_DOCX = frozenset({".docx"})

_ID_HEADERS = frozenset({"id", "field_id", "load_id", "field"})
_LABEL_HEADERS = frozenset({"label", "name", "title"})
_VALUE_HEADERS = frozenset({"value", "declared", "amount"})
_EXPECTED_HEADERS = frozenset({"expected"})
_OBSERVED_HEADERS = frozenset({"observed"})
_UNIT_HEADERS = frozenset({"unit", "ед", "ед."})


@dataclass(frozen=True)
class ExtractedTable:
    status: ExtractStatus
    rows: tuple[DeclaredCalcRow, ...]
    path_suffix: str
    reason: str


def extract_declared_rows(path: Path) -> ExtractedTable:
    """Read a declared-field table. Never parses native LIRA binaries or PDF tables."""

    suffix = path.suffix.lower()
    if suffix in _NATIVE_LIR:
        return ExtractedTable(
            status="native_lir_not_implemented",
            rows=(),
            path_suffix=suffix,
            reason="native .lir/.spr is not implemented",
        )
    if suffix in _PDF:
        return ExtractedTable(
            status="pdf_fragile",
            rows=(),
            path_suffix=suffix,
            reason="PDF table compare remains fragile; use xlsx/docx",
        )
    if not path.is_file():
        return ExtractedTable(
            status="unreadable",
            rows=(),
            path_suffix=suffix,
            reason="path is not a file",
        )
    if suffix in _XLSX:
        return _extract_xlsx(path, suffix)
    if suffix in _DOCX:
        return _extract_docx(path, suffix)
    return ExtractedTable(
        status="unsupported",
        rows=(),
        path_suffix=suffix,
        reason=f"suffix {suffix} is not an office declared-table source",
    )


def compare_office_tables(calc_path: Path, bim_path: Path) -> TableCompareResult:
    """Two-file office compare. Callers must not treat MATCH as solver OK."""

    calc = extract_declared_rows(calc_path)
    bim = extract_declared_rows(bim_path)
    return compare_declared_tables(calc.rows, bim.rows)


def _norm_header(raw: object) -> str:
    return str(raw or "").strip().casefold()


def _header_map(cells: list[str]) -> dict[str, int] | None:
    mapping: dict[str, int] = {}
    for index, cell in enumerate(cells):
        token = _norm_header(cell)
        if token in _ID_HEADERS:
            mapping["id"] = index
        elif token in _LABEL_HEADERS:
            mapping["label"] = index
        elif token in _VALUE_HEADERS:
            mapping["value"] = index
        elif token in _EXPECTED_HEADERS:
            mapping["expected"] = index
        elif token in _OBSERVED_HEADERS:
            mapping["observed"] = index
        elif token in _UNIT_HEADERS:
            mapping["unit"] = index
    if "id" not in mapping:
        return None
    if "value" not in mapping and "expected" not in mapping:
        return None
    return mapping


def _row_from_cells(cells: list[str], mapping: dict[str, int]) -> DeclaredCalcRow | None:
    field_id = cells[mapping["id"]].strip() if mapping["id"] < len(cells) else ""
    if not field_id:
        return None
    label = ""
    if "label" in mapping and mapping["label"] < len(cells):
        label = cells[mapping["label"]].strip()
    unit = ""
    if "unit" in mapping and mapping["unit"] < len(cells):
        unit = cells[mapping["unit"]].strip()
    value = ""
    if "value" in mapping and mapping["value"] < len(cells):
        value = cells[mapping["value"]].strip()
    elif "expected" in mapping and mapping["expected"] < len(cells):
        value = cells[mapping["expected"]].strip()
    if not value:
        return None
    return DeclaredCalcRow(field_id=field_id, label=label, value=value, unit=unit)


def _extract_xlsx(path: Path, suffix: str) -> ExtractedTable:
    try:
        from openpyxl import load_workbook
    except ModuleNotFoundError:
        return ExtractedTable(
            status="unreadable",
            rows=(),
            path_suffix=suffix,
            reason="openpyxl is required for .xlsx declared tables",
        )
    workbook = load_workbook(path, read_only=True, data_only=True)
    try:
        sheet = workbook[workbook.sheetnames[0]]
        matrix: list[list[str]] = []
        for row in sheet.iter_rows(values_only=True):
            matrix.append(["" if cell is None else str(cell).strip() for cell in row])
    finally:
        workbook.close()
    return _rows_from_matrix(matrix, suffix)


def _extract_docx(path: Path, suffix: str) -> ExtractedTable:
    try:
        from docx import Document
    except ModuleNotFoundError:
        return ExtractedTable(
            status="unreadable",
            rows=(),
            path_suffix=suffix,
            reason="python-docx is required for .docx declared tables",
        )
    document = Document(str(path))
    if not document.tables:
        return ExtractedTable(
            status="empty",
            rows=(),
            path_suffix=suffix,
            reason="docx has no tables",
        )
    table = document.tables[0]
    matrix = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    return _rows_from_matrix(matrix, suffix)


def _rows_from_matrix(matrix: list[list[str]], suffix: str) -> ExtractedTable:
    mapping: dict[str, int] | None = None
    start = 0
    for index, row in enumerate(matrix):
        found = _header_map(row)
        if found is not None:
            mapping = found
            start = index + 1
            break
    if mapping is None:
        return ExtractedTable(
            status="empty",
            rows=(),
            path_suffix=suffix,
            reason="no id+value/expected header row",
        )
    rows: list[DeclaredCalcRow] = []
    for row in matrix[start:]:
        parsed = _row_from_cells(row, mapping)
        if parsed is not None:
            rows.append(parsed)
    if not rows:
        return ExtractedTable(
            status="empty",
            rows=(),
            path_suffix=suffix,
            reason="header found but no data rows",
        )
    return ExtractedTable(status="ok", rows=tuple(rows), path_suffix=suffix, reason="ok")
